import os

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor, Cm, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
except ImportError:
    raise ImportError(
        "python-docx is required. Install it with:\n  pip install python-docx"
    )

# ─────────────────────────────────────────────────────────────────────────────
# Configuration
# ─────────────────────────────────────────────────────────────────────────────

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT_FILE = os.path.join(OUTPUT_DIR, "College_Student_Module_Documentation.docx")

PRIMARY_COLOR = RGBColor(0x1A, 0x56, 0xDB)   
ACCENT_COLOR = RGBColor(0x0D, 0x94, 0x88)    
DARK_TEXT = RGBColor(0x1F, 0x1F, 0x1F)
LIGHT_TEXT = RGBColor(0x4A, 0x4A, 0x4A)
TABLE_HEADER_BG = "1A56DB"
TABLE_ALT_ROW = "F0F4FF"
WHITE = "FFFFFF"

# ─────────────────────────────────────────────────────────────────────────────
# Helper Functions
# ─────────────────────────────────────────────────────────────────────────────

def set_cell_shading(cell, color_hex: str):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    cell._element.get_or_add_tcPr().append(shading_elm)

def add_formatted_paragraph(doc, text, style="Normal", bold=False,
                            font_size=11, color=None, alignment=None,
                            space_before=0, space_after=6):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(font_size)
    run.font.name = "Calibri"
    if color:
        run.font.color.rgb = color
    if alignment is not None:
        para.alignment = alignment
    pf = para.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    return para

def add_section_heading(doc, number, title):
    sep = doc.add_paragraph()
    sep_run = sep.add_run("=" * 60)
    sep_run.font.size = Pt(10)
    sep_run.font.color.rgb = PRIMARY_COLOR
    sep.paragraph_format.space_before = Pt(18)
    sep.paragraph_format.space_after = Pt(4)

    heading = doc.add_paragraph()
    run = heading.add_run(f"{number}. {title.upper()}")
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = "Calibri"
    run.font.color.rgb = PRIMARY_COLOR
    heading.paragraph_format.space_before = Pt(2)
    heading.paragraph_format.space_after = Pt(4)

    sep2 = doc.add_paragraph()
    sep2_run = sep2.add_run("=" * 60)
    sep2_run.font.size = Pt(10)
    sep2_run.font.color.rgb = PRIMARY_COLOR
    sep2.paragraph_format.space_before = Pt(0)
    sep2.paragraph_format.space_after = Pt(12)

def add_sub_heading(doc, title, font_size=13):
    para = doc.add_paragraph()
    run = para.add_run(title)
    run.bold = True
    run.font.size = Pt(font_size)
    run.font.name = "Calibri"
    run.font.color.rgb = ACCENT_COLOR
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(6)
    return para

def add_bullet_point(doc, text, indent_level=0):
    para = doc.add_paragraph(style="List Bullet")
    para.clear()
    run = para.add_run(text)
    run.font.size = Pt(11)
    run.font.name = "Calibri"
    run.font.color.rgb = DARK_TEXT
    para.paragraph_format.space_after = Pt(3)
    if indent_level:
        para.paragraph_format.left_indent = Inches(0.25 * indent_level)
    return para

def add_score_table(doc, rows, header_labels=("Criteria", "Score")):
    table = doc.add_table(rows=1 + len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    for row_obj in table.rows:
        row_obj.cells[0].width = Inches(4.0)
        row_obj.cells[1].width = Inches(1.5)

    for idx, label in enumerate(header_labels):
        cell = table.rows[0].cells[idx]
        cell.text = ""
        run = cell.paragraphs[0].add_run(label)
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, TABLE_HEADER_BG)

    for i, (criteria, score) in enumerate(rows):
        row = table.rows[i + 1]
        c0 = row.cells[0]
        c0.text = ""
        run0 = c0.paragraphs[0].add_run(criteria)
        run0.font.size = Pt(11)
        run0.font.name = "Calibri"

        c1 = row.cells[1]
        c1.text = ""
        run1 = c1.paragraphs[0].add_run(score)
        run1.font.size = Pt(11)
        run1.font.name = "Calibri"
        run1.bold = True
        c1.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        if i % 2 == 0:
            set_cell_shading(c0, TABLE_ALT_ROW)
            set_cell_shading(c1, TABLE_ALT_ROW)

    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="6" w:space="0" w:color="CCCCCC"/>'
        '  <w:left w:val="single" w:sz="6" w:space="0" w:color="CCCCCC"/>'
        '  <w:bottom w:val="single" w:sz="6" w:space="0" w:color="CCCCCC"/>'
        '  <w:right w:val="single" w:sz="6" w:space="0" w:color="CCCCCC"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="DDDDDD"/>'
        '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="DDDDDD"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)
    return table

def add_three_column_table(doc, rows, header_labels=("Required Skill", "Student Level", "Status")):
    table = doc.add_table(rows=1 + len(rows), cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    for row_obj in table.rows:
        row_obj.cells[0].width = Inches(2.5)
        row_obj.cells[1].width = Inches(1.5)
        row_obj.cells[2].width = Inches(1.5)

    for idx, label in enumerate(header_labels):
        cell = table.rows[0].cells[idx]
        cell.text = ""
        run = cell.paragraphs[0].add_run(label)
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, TABLE_HEADER_BG)

    for i, (col1, col2, col3) in enumerate(rows):
        row = table.rows[i + 1]
        
        for j, val in enumerate([col1, col2, col3]):
            c = row.cells[j]
            c.text = ""
            run = c.paragraphs[0].add_run(val)
            run.font.size = Pt(11)
            run.font.name = "Calibri"
            if j > 0:
                c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            if i % 2 == 0:
                set_cell_shading(c, TABLE_ALT_ROW)

    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="6" w:space="0" w:color="CCCCCC"/>'
        '  <w:left w:val="single" w:sz="6" w:space="0" w:color="CCCCCC"/>'
        '  <w:bottom w:val="single" w:sz="6" w:space="0" w:color="CCCCCC"/>'
        '  <w:right w:val="single" w:sz="6" w:space="0" w:color="CCCCCC"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="DDDDDD"/>'
        '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="DDDDDD"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)
    return table

def add_five_column_table(doc, rows, header_labels=("Company", "Role", "Location", "Required Skills", "Eligibility")):
    table = doc.add_table(rows=1 + len(rows), cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    widths = [1.0, 1.5, 1.0, 1.5, 1.5]
    for row_obj in table.rows:
        for idx, width in enumerate(widths):
            row_obj.cells[idx].width = Inches(width)

    for idx, label in enumerate(header_labels):
        cell = table.rows[0].cells[idx]
        cell.text = ""
        run = cell.paragraphs[0].add_run(label)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, TABLE_HEADER_BG)

    for i, row_data in enumerate(rows):
        row = table.rows[i + 1]
        
        for j, val in enumerate(row_data):
            c = row.cells[j]
            c.text = ""
            run = c.paragraphs[0].add_run(val)
            run.font.size = Pt(10)
            run.font.name = "Calibri"
            if j > 0:
                c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            if i % 2 == 0:
                set_cell_shading(c, TABLE_ALT_ROW)

    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="6" w:space="0" w:color="CCCCCC"/>'
        '  <w:left w:val="single" w:sz="6" w:space="0" w:color="CCCCCC"/>'
        '  <w:bottom w:val="single" w:sz="6" w:space="0" w:color="CCCCCC"/>'
        '  <w:right w:val="single" w:sz="6" w:space="0" w:color="CCCCCC"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="DDDDDD"/>'
        '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="DDDDDD"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)
    return table

def add_feedback_box(doc, feedback_text):
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Inches(0.5)
    para.paragraph_format.right_indent = Inches(0.5)
    para.paragraph_format.space_before = Pt(8)
    para.paragraph_format.space_after = Pt(8)

    run = para.add_run(f'"{feedback_text}"')
    run.italic = True
    run.font.size = Pt(11)
    run.font.name = "Calibri"
    run.font.color.rgb = LIGHT_TEXT
    return para

# ─────────────────────────────────────────────────────────────────────────────
# Document Generation
# ─────────────────────────────────────────────────────────────────────────────

def generate_document():
    doc = Document()

    # ── Page setup ──────────────────────────────────────────────────────
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = DARK_TEXT

    # ════════════════════════════════════════════════════════════════════
    # TITLE PAGE
    # ════════════════════════════════════════════════════════════════════

    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("TalentSphere Elevate")
    run.bold = True
    run.font.size = Pt(28)
    run.font.name = "Calibri"
    run.font.color.rgb = PRIMARY_COLOR

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("College Student Module — Feature Documentation")
    run.font.size = Pt(16)
    run.font.name = "Calibri"
    run.font.color.rgb = ACCENT_COLOR

    doc.add_paragraph()

    version = doc.add_paragraph()
    version.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = version.add_run("Version 3.0 — Updated Implementation")
    run.font.size = Pt(12)
    run.font.name = "Calibri"
    run.font.color.rgb = LIGHT_TEXT

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════
    # TABLE OF CONTENTS
    # ════════════════════════════════════════════════════════════════════

    toc_title = doc.add_paragraph()
    run = toc_title.add_run("TABLE OF CONTENTS")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = PRIMARY_COLOR
    toc_title.paragraph_format.space_after = Pt(12)

    toc_items = [
        "1. Student Profile",
        "2. ATS Friendly Resume Builder",
        "3. AI Resume Review",
        "4. Coding Practice Tests",
        "5. Placement Readiness Score",
        "6. Mock Interview System",
        "7. Skill Gap Analysis",
        "8. Internship Recommendations",
        "9. Job Matching Engine",
        "10. AI Output Example",
        "11. College Dashboard",
    ]
    for item in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(item)
        run.font.size = Pt(12)
        run.font.name = "Calibri"
        run.font.color.rgb = DARK_TEXT
        p.paragraph_format.space_after = Pt(4)

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════
    # 1. STUDENT PROFILE
    # ════════════════════════════════════════════════════════════════════

    add_section_heading(doc, 1, "Student Profile")

    add_formatted_paragraph(
        doc,
        "The Student Profile captures comprehensive academic, personal, and career information for each college student. This information is used by the AI engine to personalize recommendations, ATS analysis, internships, career guidance, reports, and job matching.",
        font_size=11,
        color=DARK_TEXT,
        space_after=8,
    )

    add_sub_heading(doc, "Resume Upload")
    add_bullet_point(doc, "Resume Upload (PDF/DOCX)")

    add_sub_heading(doc, "Personal Information")
    for item in [
        "Full Name",
        "Date of Birth",
        "Gender",
        "Email Address",
        "Mobile Number",
        "State",
        "City",
    ]:
        add_bullet_point(doc, item)

    add_sub_heading(doc, "Academic Information")
    for item in [
        "College Name",
        "University Name",
        "Degree",
        "Department / Branch",
        "Year of Study",
        "Semester",
        "CGPA",
        "Percentage",
    ]:
        add_bullet_point(doc, item)

    add_sub_heading(doc, "Technical Skills")
    for item in [
        "Programming Languages",
        "Frameworks",
        "Databases",
        "Development Tools",
    ]:
        add_bullet_point(doc, item)

    add_sub_heading(doc, "Career Information")
    for item in [
        "Interested Job Roles",
        "Preferred Industry",
        "Preferred Work Location",
        "Internship / Full-time Preference",
    ]:
        add_bullet_point(doc, item)

    add_sub_heading(doc, "Career Goals")
    for item in [
        "Short-Term Goal",
        "Long-Term Goal",
        "Dream Role",
        "Target Company",
    ]:
        add_bullet_point(doc, item)

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════
    # 2. ATS FRIENDLY RESUME BUILDER
    # ════════════════════════════════════════════════════════════════════

    add_section_heading(doc, 2, "ATS Friendly Resume Builder")

    add_sub_heading(doc, "Features Included")
    for item in [
        "ATS Resume Templates",
        "Resume Builder",
        "Resume Preview",
        "ATS Resume Analysis",
        "Resume Score",
        "PDF Download",
    ]:
        add_bullet_point(doc, item)

    add_sub_heading(doc, "Resume Sections")
    for item in [
        "Personal Details",
        "Professional Links (GitHub, LinkedIn)",
        "Education",
        "Skills",
        "Projects",
        "Experience",
        "Certifications",
        "Achievements",
        "Languages",
        "Career Objective",
    ]:
        add_bullet_point(doc, item)

    add_sub_heading(doc, "AI Suggestions")
    for item in [
        "ATS keyword optimization",
        "Resume improvements",
        "Missing technical keywords",
        "Better project descriptions",
        "Resume formatting suggestions",
        "Strong action verbs",
        "Career objective improvements",
    ]:
        add_bullet_point(doc, item)

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════
    # 3. AI RESUME REVIEW
    # ════════════════════════════════════════════════════════════════════

    add_section_heading(doc, 3, "AI Resume Review")

    add_sub_heading(doc, "AI Evaluates")
    for item in [
        "ATS Compatibility",
        "Resume Score",
        "Keyword Optimization",
        "Technical Skills",
        "Education",
        "Projects",
        "Experience",
        "Resume Formatting",
    ]:
        add_bullet_point(doc, item)

    add_sub_heading(doc, "Generated Output")
    for item in [
        "Overall Resume Score",
        "AI Feedback",
        "Missing Keywords",
        "Resume Recommendations",
    ]:
        add_bullet_point(doc, item)

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════
    # 4. CODING PRACTICE TESTS
    # ════════════════════════════════════════════════════════════════════

    add_section_heading(doc, 4, "Coding Practice Tests")

    add_sub_heading(doc, "Topics")
    for item in [
        "Python",
        "Java",
        "C/C++",
        "SQL",
        "DBMS",
        "Data Structures",
        "Algorithms",
        "HTML",
        "CSS",
        "JavaScript",
        "React",
        "Node.js",
        "Operating Systems",
        "Computer Networks",
        "Aptitude",
        "Logical Reasoning",
    ]:
        add_bullet_point(doc, item)

    add_sub_heading(doc, "Generated Output")
    for item in [
        "Test Scores",
        "Performance Analysis",
        "Topic-wise Progress",
        "AI Coding Suggestions",
    ]:
        add_bullet_point(doc, item)

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════
    # 5. PLACEMENT READINESS SCORE
    # ════════════════════════════════════════════════════════════════════

    add_section_heading(doc, 5, "Placement Readiness Score")

    add_sub_heading(doc, "Calculated Using")
    for item in [
        "Resume Score",
        "ATS Score",
        "Coding Score",
        "Technical Skills",
        "Communication Skills",
        "Projects",
        "Internship Experience",
        "Certifications",
    ]:
        add_bullet_point(doc, item)

    add_sub_heading(doc, "Display")
    for item in [
        "Placement Readiness Formula",
        "Final Placement Score",
        "AI Placement Recommendations",
    ]:
        add_bullet_point(doc, item)

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════
    # 6. MOCK INTERVIEW SYSTEM
    # ════════════════════════════════════════════════════════════════════

    add_section_heading(doc, 6, "Mock Interview System")

    add_sub_heading(doc, "Included Components")
    for item in [
        "Technical Interview",
        "HR Interview",
        "Communication Evaluation",
        "Confidence Evaluation",
        "Problem Solving",
        "AI Feedback",
        "AI Suggestions",
        "Overall Interview Score",
    ]:
        add_bullet_point(doc, item)

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════
    # 7. SKILL GAP ANALYSIS
    # ════════════════════════════════════════════════════════════════════

    add_section_heading(doc, 7, "Skill Gap Analysis")

    add_sub_heading(doc, "Display")
    for item in [
        "Current Skills",
        "Missing Skills",
        "Skill Match Percentage",
        "Required Skills",
        "AI Learning Recommendations",
        "Suggested Learning Resources",
        "Personalized Learning Roadmap",
    ]:
        add_bullet_point(doc, item)

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════
    # 8. INTERNSHIP RECOMMENDATIONS
    # ════════════════════════════════════════════════════════════════════

    add_section_heading(doc, 8, "Internship Recommendations")

    add_sub_heading(doc, "Display")
    for item in [
        "Company",
        "Internship Role",
        "Location",
        "Required Skills",
        "Eligibility",
        "Internship Platform",
    ]:
        add_bullet_point(doc, item)

    add_sub_heading(doc, "AI Recommendations")
    add_bullet_point(doc, "Generate AI Internship Recommendations.")

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════
    # 9. JOB MATCHING ENGINE
    # ════════════════════════════════════════════════════════════════════

    add_section_heading(doc, 9, "Job Matching Engine")

    add_sub_heading(doc, "Job Match % Display For")
    for item in [
        "Software Developer",
        "Backend Developer",
        "Full Stack Developer",
        "Python Developer",
        "AI Engineer",
        "Data Analyst",
        "Cloud Engineer",
        "QA Engineer",
        "DevOps Engineer",
        "Cyber Security Analyst",
    ]:
        add_bullet_point(doc, item)

    add_sub_heading(doc, "AI Matching Factors Explained")
    for item in [
        "Skills",
        "Resume Score",
        "Coding Performance",
        "Projects",
        "CGPA",
        "Career Interests",
    ]:
        add_bullet_point(doc, item)

    add_sub_heading(doc, "Career Recommendations")
    add_bullet_point(doc, "Include Personalized Career Recommendations.")

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════
    # 10. AI OUTPUT EXAMPLE
    # ════════════════════════════════════════════════════════════════════

    add_section_heading(doc, 10, "AI Output Example")

    add_sub_heading(doc, "Professional AI Report Shows")
    for item in [
        "Resume Score",
        "ATS Score",
        "Coding Score",
        "Placement Readiness",
        "Technical Skills",
        "Job Match",
        "Missing Skills",
        "Strong Skills",
        "Suggested Certifications",
        "Internship Recommendations",
        "Career Path",
        "30-Day Action Plan",
        "Final AI Recommendation",
    ]:
        add_bullet_point(doc, item)

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════
    # 11. COLLEGE DASHBOARD
    # ════════════════════════════════════════════════════════════════════

    add_section_heading(doc, 11, "College Dashboard")

    add_sub_heading(doc, "Dashboard Includes")
    for item in [
        "Placement Readiness",
        "ATS Resume Score",
        "Technical Skills",
        "Coding Performance",
        "Job Match",
        "Resume Completion",
        "Assessments Completed",
        "Active Roadmaps",
        "Recommended Job Role",
        "Current Skill Progress",
        "AI Dashboard Insights",
        "Monthly Focus Areas",
    ]:
        add_bullet_point(doc, item)

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════
    # EXPECTED OUTCOME
    # ════════════════════════════════════════════════════════════════════

    outcome_heading = doc.add_paragraph()
    run = outcome_heading.add_run("EXPECTED OUTCOME")
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = "Calibri"
    run.font.color.rgb = PRIMARY_COLOR
    outcome_heading.paragraph_format.space_before = Pt(14)
    outcome_heading.paragraph_format.space_after = Pt(8)

    add_formatted_paragraph(
        doc,
        "The College Student Module provides:",
        font_size=11,
        color=DARK_TEXT,
        space_after=8,
    )

    for item in [
        "AI Resume Builder",
        "ATS Resume Analysis",
        "Coding Practice",
        "Skill Gap Analysis",
        "Placement Readiness",
        "Mock Interviews",
        "Internship Recommendations",
        "Job Matching",
        "Personalized Learning Roadmaps",
        "Career Guidance",
        "Industry Readiness",
    ]:
        add_bullet_point(doc, item)

    # ════════════════════════════════════════════════════════════════════
    # FINAL PROJECT STATEMENT
    # ════════════════════════════════════════════════════════════════════

    final_heading = doc.add_paragraph()
    run = final_heading.add_run("FINAL PROJECT STATEMENT")
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = "Calibri"
    run.font.color.rgb = PRIMARY_COLOR
    final_heading.paragraph_format.space_before = Pt(14)
    final_heading.paragraph_format.space_after = Pt(8)

    add_feedback_box(
        doc,
        "The College Student Module integrates AI-powered Resume Building, ATS Analysis, "
        "Resume Review, Coding Assessment, Skill Gap Analysis, Placement Readiness, Mock Interviews, "
        "Internship Recommendation, Job Matching, Personalized Learning Roadmaps, and Career Guidance "
        "into one intelligent platform that prepares students for placements and professional careers."
    )

    # ── Add page numbers ────────────────────────────────────────────────
    for sec in doc.sections:
        footer = sec.footer
        footer.is_linked_to_previous = False
        footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        run = footer_para.add_run()
        fld_char_begin = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run._element.append(fld_char_begin)
        
        run2 = footer_para.add_run()
        instr = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
        run2._element.append(instr)
        
        run3 = footer_para.add_run()
        fld_char_end = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run3._element.append(fld_char_end)

    # ── Save ────────────────────────────────────────────────────────────
    doc.save(OUTPUT_FILE)
    print(f"✅ Document generated successfully!")
    print(f"📄 Output: {OUTPUT_FILE}")


if __name__ == "__main__":
    generate_document()
